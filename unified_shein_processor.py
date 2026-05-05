import os
import threading
import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from docx.shared import Cm, Inches
from docx.enum.table import WD_ROW_HEIGHT_RULE
import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
from datetime import datetime
import logging
import platform
import subprocess
import queue
import shutil
from tkinterdnd2 import TkinterDnD, DND_FILES

try:
    from docx2pdf import convert
except ImportError:
    convert = None

# ============================================================================
# LOGGING SETUP
# ============================================================================
class QueueHandler(logging.Handler):
    """Log handler sending messages to a queue."""
    def __init__(self, log_queue):
        super().__init__()
        self.log_queue = log_queue

    def emit(self, record):
        self.log_queue.put(self.format(record))

# ============================================================================
# UTILITIES
# ============================================================================
def open_folder(path: str):
    """Abre la carpeta en el explorador según el SO."""
    try:
        if platform.system() == "Windows":
            os.startfile(path)
        elif platform.system() == "Darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
    except Exception as e:
        print(f"Error opening folder: {e}")

def remove_even_pages(pdf_in: str, pdf_out: str):
    """Genera un PDF con solo páginas impares (usado p.ej. en TikTok laser)."""
    try:
        if not os.path.exists(pdf_in):
            return False, f"⚠️ PDF de origen no existe: {pdf_in}"
        src = fitz.open(pdf_in)
        dst = fitz.open()
        for i in range(src.page_count):
            if (i + 1) % 2 == 1:
                dst.insert_pdf(src, from_page=i, to_page=i)
        dst.save(pdf_out)
        dst.close()
        src.close()
        return True, f"✅ PDF (impares extraídos) generado: {os.path.basename(pdf_out)}"
    except Exception as e:
        return False, f"❌ Error al filtrar impares: {e}"

# ============================================================================
# CORE PROCESSING
# ============================================================================
def process_pdf_for_day(pdf_path, brand, directory, today, logger, shortcuts_dict, day_label=""):
    """
    Procesa un único PDF. 
    directory: la carpeta en la que deben guardarse todos los resultados de este día.
    shortcuts_dict: diccionario con rutas 'imprimir' y 'correo' para auto-copia.
    Devuelve: (booleano_exito, cantidad_de_guias_procesadas)
    """
    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        logger.error(f"❌ Error abriendo PDF ({day_label}): {e}")
        return False, 0

    total = doc.page_count
    logger.info(f"📄 Total páginas [{day_label or 'Único'}]: {total}")

    day_suffix = f" {day_label}" if day_label else ""

    marketplace = "tiktok" if brand == "TikTok" else "shein"

    # --- A) IMPARES -> PDF TÉRMICO ---
    odd_pdf = fitz.open()
    odd_indices = []
    for i in range(total):
        if (i + 1) % 2 == 1:
            odd_pdf.insert_pdf(doc, from_page=i, to_page=i)
            odd_indices.append(i + 1)
    
    odd_name = f"{brand} Guias {marketplace} {today}{day_suffix} - impresora termica.pdf"
    odd_path = os.path.join(directory, odd_name)
    try:
        odd_pdf.save(odd_path)
        shutil.copy2(odd_path, os.path.join(shortcuts_dict["imprimir"], odd_name))
        logger.info(f"✅ Térmico guardado ({odd_name}).")
    except Exception as e:
        logger.error(f"❌ Error guardando térmico: {e}")
    finally:
        odd_pdf.close()

    # --- B) PARES -> DOCX (Imágenes) ---
    even_indices = [i + 1 for i in range(total) if (i + 1) % 2 == 0]
    images = []
    scale = 2  # Zoom para calidad
    
    # Extraer imágenes pares
    for idx in even_indices:
        try:
            page = doc.load_page(idx - 1)
            pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            if img.getbbox():
                images.append((idx, img))
        except Exception as e:
            logger.warning(f"⚠️ Error leyendo página {idx}: {e}")

    # Configurar Word
    doc_word = Document()
    for section in doc_word.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)

    section = doc_word.sections[0]
    usable_w_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    usable_h_cm = section.page_height.cm - section.top_margin.cm - section.bottom_margin.cm

    if brand == "TikTok":
        per_page = 2
        rows, cols = 2, 1
        picture_width = Cm(usable_w_cm * 0.98)
        out_suffix = "tiktok laser"
    else:
        per_page = 4
        rows, cols = 2, 2
        picture_width = Cm(7.0)
        out_suffix = "impresora laser"

    count = 0
    table = None
    temps_to_delete = []
    
    for idx, img in images:
        if count % per_page == 0:
            table = doc_word.add_table(rows=rows, cols=cols)
            table.autofit = False
            table.allow_autofit = False
            
            if brand == "TikTok":
                half_h_cm = usable_h_cm / 2.0
                for r in range(2):
                    table.rows[r].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                    table.rows[r].height = Cm(half_h_cm)

        pos = count % per_page
        row_i = pos // cols
        col_i = pos % cols
        cell = table.cell(row_i, col_i)

        temp_png = os.path.join(directory, f"temp_{brand}_{day_label}_{idx}.png")
        img.save(temp_png, quality=100)
        temps_to_delete.append(temp_png)
        
        try:
            run = cell.paragraphs[0].add_run()
            run.add_picture(temp_png, width=picture_width)
            count += 1
            if count % per_page == 0:
                doc_word.add_page_break()
        except Exception as e:
            logger.error(f"❌ Error insertando img {idx} en Word: {e}")

    docx_name = f"{brand} Guias {marketplace} {today}{day_suffix} - {out_suffix}.docx"
    docx_path = os.path.join(directory, docx_name)
    try:
        doc_word.save(docx_path)
        shutil.copy2(docx_path, os.path.join(shortcuts_dict["correo"], docx_name))
        logger.info(f"✅ DOCX láser guardado ({docx_name}).")
    except Exception as e:
        logger.error(f"❌ Error guardando DOCX: {e}")

    # Limpiar PNGs temporales
    for tmp in temps_to_delete:
        try:
            os.remove(tmp)
        except:
            pass

    # --- C) DOCX -> PDF ---
    produced_pdf_path = None
    if convert:
        pdf_name = f"{brand} Guias {marketplace} {today}{day_suffix} - {out_suffix}.pdf"
        produced_pdf_path = os.path.join(directory, pdf_name)
        try:
            logger.info(f"⏳ Convirtiendo a PDF (Tomará un momento)...")
            convert(docx_path, produced_pdf_path)
            shutil.copy2(produced_pdf_path, os.path.join(shortcuts_dict["correo"], pdf_name))
            logger.info("✅ PDF Láser generado.")
        except Exception as e:
            logger.error(f"⚠️ Falló conversión DOCX->PDF: {e}")
    else:
        logger.warning("⚠️ docx2pdf no está instalado. Omitiendo conversión a PDF.")

    # --- D) TikTok Extra: Impares Laser ---
    if brand == "TikTok" and produced_pdf_path:
        laser_odds_pdf_name = f"{brand} Guias {marketplace} {today}{day_suffix} - laser (solo impares).pdf"
        laser_odds_pdf = os.path.join(directory, laser_odds_pdf_name)
        ok, msg = remove_even_pages(produced_pdf_path, laser_odds_pdf)
        if ok:
            shutil.copy2(laser_odds_pdf, os.path.join(shortcuts_dict["correo"], laser_odds_pdf_name))
            logger.info(msg)
        else:
            logger.error(msg)
    
    doc.close()
    return True, len(images)

def run_processing_pipeline(brand, mode, files_dict, logger, clear_shortcuts=False):
    """
    Función envoltorio para gestionar las carpetas y los diferentes modos (Único, Semana, Puente).
    """
    logger.info("==========================================")
    logger.info(f"▶ INICIANDO PROCESO | MODO: {mode}")
    logger.info(f"🔖 Marca: {brand}")
    
    today = datetime.now().strftime("%d-%m-%Y")
    desktop = os.path.expanduser("~/Desktop")
    base_shein = os.path.join(desktop, "guias-shein")
    
    imprimir_dir = os.path.join(base_shein, "imprimir")
    correo_dir = os.path.join(base_shein, "correo")
    
    if clear_shortcuts:
        logger.info("🗑 Vaciando carpetas cortafuegos (imprimir, correo)...")
        for d in [imprimir_dir, correo_dir]:
            if os.path.exists(d):
                for filename in os.listdir(d):
                    file_path = os.path.join(d, filename)
                    try:
                        if os.path.isfile(file_path):
                            os.unlink(file_path)
                        elif os.path.isdir(file_path):
                            shutil.rmtree(file_path)
                    except Exception as e:
                        logger.warning(f"⚠️ No se pudo limpiar {file_path}: {e}")
                        
    os.makedirs(imprimir_dir, exist_ok=True)
    os.makedirs(correo_dir, exist_ok=True)
    shortcuts_dict = {"imprimir": imprimir_dir, "correo": correo_dir}
    
    # Crear carpeta principal para el día de hoy
    safe_brand = brand.replace(" ", "_")
    day_folder_root = os.path.join(base_shein, today, safe_brand)
    os.makedirs(day_folder_root, exist_ok=True)
    
    logger.info(f"📂 Guardando maestras en: {day_folder_root}")

    reporte_guias = {}

    if mode == "Día Único":
        pdf_path = files_dict.get("Único")
        if not pdf_path:
            logger.error("❌ Archivo Único no especificado.")
            return False, day_folder_root, {}
        
        success, count = process_pdf_for_day(pdf_path, brand, day_folder_root, today, logger, shortcuts_dict, day_label="")
        if success:
            reporte_guias["Único"] = count

    else:
        # Fin de Semana / Puente
        days_to_process = ["Viernes", "Sabado", "Domingo"]
        if mode == "Puente (4 Días)":
            days_to_process.append("Lunes")
            
        for day in days_to_process:
            pdf_path = files_dict.get(day)
            if not pdf_path:
                logger.error(f"❌ Archivo para '{day}' no proporcionado.")
                continue
                
            logger.info(f"\n--- 🗓 Procesando {day} ---")
            day_folder = os.path.join(day_folder_root, day)
            os.makedirs(day_folder, exist_ok=True)
            success, count = process_pdf_for_day(pdf_path, brand, day_folder, today, logger, shortcuts_dict, day_label=day)
            if success:
                reporte_guias[day] = count

    logger.info("✨ PROCESO COMPLETADO EXITOSAMENTE ✨")
    return True, day_folder_root, reporte_guias

# ============================================================================
# GUI / APPLICATION
# ============================================================================
class CustomAlert(ctk.CTkToplevel):
    def __init__(self, master, title, message, is_error=False, accent_color="#0066cc"):
        super().__init__(master)
        self.title(" " + title)
        self.geometry("450x250")
        self.resizable(False, False)
        
        self.attributes("-topmost", True)
        self.focus_force()
        self.grab_set()

        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)
        
        frame = ctk.CTkFrame(self, fg_color="transparent")
        frame.grid(row=0, column=0, sticky="nsew", padx=20, pady=20)
        frame.grid_rowconfigure(0, weight=1)
        frame.grid_columnconfigure(0, weight=1)

        icon = "❌ " if is_error else "✅ "
        lbl_msg = ctk.CTkLabel(frame, text=icon + message, font=ctk.CTkFont(size=14), wraplength=380, justify="left")
        lbl_msg.grid(row=0, column=0, pady=(10, 20), sticky="nsew")
        
        btn_color = "#e11d48" if is_error else accent_color
        btn_ok = ctk.CTkButton(frame, text="Aceptar", command=self.destroy, fg_color=btn_color, height=40, font=ctk.CTkFont(weight="bold"))
        btn_ok.grid(row=1, column=0, pady=(0, 10))

class TkinterDnD_CTk(ctk.CTk, TkinterDnD.DnDWrapper):
    """Mix-in class to enable TkinterDnD within CustomTkinter."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.TkdndVersion = TkinterDnD._require(self)

class UnifiedSheinApp(TkinterDnD_CTk):
    def __init__(self):
        super().__init__()

        # Configuración de Tkinter Theme (Clean Dark Modern)
        ctk.set_appearance_mode("Dark")
        ctk.set_default_color_theme("blue")
        
        self.title("Sistema de Procesamiento Shein")
        self.geometry("900x650")
        self.minsize(800, 600)

        # Variables de estado
        self.log_queue = queue.Queue()
        self.processing = False
        self.result_folder = None
        self.folders_cleared_this_session = False
        
        self.mode_var = ctk.StringVar(value="Día Único")
        self.brand_var = ctk.StringVar(value="Marcas y Licencias")
        self.files_selected = {}
        
        # UI Layout principal
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self.setup_sidebar()
        self.setup_main_area()

        # Variables para colores dinámicos
        self.brand_colors = {
            "Marcas y Licencias": {"primary": "#0066cc", "hover": "#0052a3", "accent": "#0066cc"}, # Azul
            "Pure and Simple": {"primary": "#28a745", "hover": "#218838", "accent": "#28a745"}, # Verde
            "TikTok": {"primary": "#ff0050", "hover": "#cc0040", "accent": "#ff0050"} # Rosa/Rojo TikTok
        }
        
        # Inicializar colores dinámicos
        self.on_brand_change(self.brand_var.get())

        # Activar Drag & Drop en la ventana principal
        self.drop_target_register(DND_FILES)
        self.dnd_bind('<<Drop>>', self.on_file_drop)

        # Iniciar polling del log
        self.after(100, self.poll_log_queue)

    def setup_sidebar(self):
        # Sidebar Frame
        self.sidebar_frame = ctk.CTkFrame(self, width=250, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(7, weight=1) # Filler

        # Título
        lbl_title = ctk.CTkLabel(self.sidebar_frame, text="Shein Processor", font=ctk.CTkFont(size=22, weight="bold"))
        lbl_title.grid(row=0, column=0, padx=20, pady=(30, 20))

        # Marca
        lbl_brand = ctk.CTkLabel(self.sidebar_frame, text="Marca", font=ctk.CTkFont(size=12, weight="bold"), text_color="gray70")
        lbl_brand.grid(row=1, column=0, padx=20, pady=(10, 0), sticky="w")
        
        self.brand_menu = ctk.CTkOptionMenu(
            self.sidebar_frame, 
            values=["Marcas y Licencias", "Pure and Simple", "TikTok"],
            variable=self.brand_var,
            font=("Roboto", 13),
            command=self.on_brand_change
        )
        self.brand_menu.grid(row=2, column=0, padx=20, pady=(5, 15), sticky="ew")

        # Modo (Tiempo)
        lbl_mode = ctk.CTkLabel(self.sidebar_frame, text="Modo de Operación", font=ctk.CTkFont(size=12, weight="bold"), text_color="gray70")
        lbl_mode.grid(row=3, column=0, padx=20, pady=(10, 0), sticky="w")

        # Usamos SegmentedButton para el modo
        self.mode_selector = ctk.CTkSegmentedButton(
            self.sidebar_frame,
            values=["Día Único", "Fin de Semana", "Puente (4 Días)"],
            variable=self.mode_var,
            command=self.on_mode_change
        )
        self.mode_selector.grid(row=4, column=0, padx=20, pady=(5, 30), sticky="ew")

        # Botón de Procesar
        self.btn_process = ctk.CTkButton(
            self.sidebar_frame, 
            text="Procesar Guías", 
            font=ctk.CTkFont(size=15, weight="bold"),
            height=45,
            command=self.start_thread,
            fg_color="#0066cc",
            hover_color="#0052a3"
        )
        self.btn_process.grid(row=5, column=0, padx=20, pady=10, sticky="ew")

        # Botón Abrir Carpeta
        self.btn_open_folder = ctk.CTkButton(
            self.sidebar_frame,
            text="Abrir Carpeta Generada",
            command=self.open_output_folder,
            fg_color="transparent",
            text_color="#0066cc",
            border_width=1,
            border_color="gray50",
            hover_color="#333333",
            state="disabled"
        )
        self.btn_open_folder.grid(row=6, column=0, padx=20, pady=10, sticky="ew")

    def setup_main_area(self):
        # Frame derecho principal
        self.main_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.main_frame.grid(row=0, column=1, padx=20, pady=20, sticky="nsew")
        self.main_frame.grid_rowconfigure(2, weight=1) # El log se expande
        self.main_frame.grid_columnconfigure(0, weight=1)

        # 1. Área de selección de archivos
        self.files_container = ctk.CTkFrame(self.main_frame, corner_radius=15, fg_color="#1c1c1e")
        self.files_container.grid(row=0, column=0, sticky="ew", pady=(0, 20))
        self.files_container.grid_columnconfigure(0, weight=1)
        
        lbl_files = ctk.CTkLabel(self.files_container, text="Selección de Archivos PDF", font=ctk.CTkFont(size=16, weight="bold"))
        lbl_files.grid(row=0, column=0, pady=(15, 5), padx=20, sticky="w")

        # Sub-frame dinámico para los campos de archivo
        self.dynamic_fields_frame = ctk.CTkFrame(self.files_container, fg_color="transparent")
        self.dynamic_fields_frame.grid(row=1, column=0, sticky="ew", padx=20, pady=(0, 20))
        self.dynamic_fields_frame.grid_columnconfigure(1, weight=1)

        self.path_entries = {}
        self.build_file_fields() # Construir inicial (Día Único)

        # 2. Área de Progreso y Título Logs
        self.progress_bar = ctk.CTkProgressBar(self.main_frame, mode="indeterminate", height=6)
        self.progress_bar.grid(row=1, column=0, sticky="ew", padx=20, pady=(0, 10))
        self.progress_bar.set(0) # Visualmente apagado hasta que inicie
        
        lbl_log = ctk.CTkLabel(self.main_frame, text="Consola de Operaciones", font=ctk.CTkFont(size=14, weight="bold"))
        lbl_log.grid(row=2, column=0, sticky="w", padx=20, pady=(0, 5))

        # 3. Log Box
        self.log_box = ctk.CTkTextbox(
            self.main_frame, 
            corner_radius=10,
            fg_color="#0d0d0d", 
            text_color="#e0e0e0",
            font=("Consolas", 12),
            state="disabled"
        )
        self.log_box.grid(row=3, column=0, padx=20, sticky="nsew")
        
        # Etiquetas de color para el log
        self.log_box.tag_config("success", foreground="#a6e3a1") # Soft Green
        self.log_box.tag_config("error", foreground="#f38ba8")   # Red
        self.log_box.tag_config("warning", foreground="#f9e2af") # Yellow
        self.log_box.tag_config("header", foreground="#89b4fa")  # Blue
        self.log_box.tag_config("normal", foreground="#e0e0e0")  # Default

    def build_file_fields(self):
        # Limpiar
        for widget in self.dynamic_fields_frame.winfo_children():
            widget.destroy()
        
        self.files_selected.clear()
        self.path_entries.clear()

        mode = self.mode_var.get()
        if mode == "Día Único":
            keys = ["Único"]
        elif mode == "Fin de Semana":
            keys = ["Viernes", "Sabado", "Domingo"]
        else:
            keys = ["Viernes", "Sabado", "Domingo", "Lunes"]

        for i, key in enumerate(keys):
            lbl = ctk.CTkLabel(self.dynamic_fields_frame, text=f"{key}:", width=80, anchor="e", font=ctk.CTkFont(weight="bold"))
            lbl.grid(row=i, column=0, padx=(0, 10), pady=10, sticky="e")

            entry = ctk.CTkEntry(self.dynamic_fields_frame, placeholder_text="Seleccione un archivo...", state="disabled", fg_color="#2c2c2e", border_color="#3a3a3c")
            entry.grid(row=i, column=1, padx=(0, 10), pady=10, sticky="ew")
            self.path_entries[key] = entry

            btn = ctk.CTkButton(
                self.dynamic_fields_frame, 
                text="Examinar", 
                width=100, 
                command=lambda k=key: self.select_file(k),
                fg_color="#3a3a3c",
                hover_color="#48484a"
            )
            btn.grid(row=i, column=2, pady=10)

    def on_mode_change(self, value):
        self.build_file_fields()

    def on_brand_change(self, value):
        colors = self.brand_colors.get(value, self.brand_colors["Marcas y Licencias"])
        
        self.btn_process.configure(fg_color=colors["primary"], hover_color=colors["hover"])
        self.mode_selector.configure(selected_color=colors["primary"], selected_hover_color=colors["hover"])
        self.btn_open_folder.configure(text_color=colors["accent"])
        self.progress_bar.configure(progress_color=colors["accent"])
        
        # Limpiar seleccion al cambiar marca
        self.files_selected.clear()
        for entry in self.path_entries.values():
            entry.configure(state="normal")
            entry.delete(0, "end")
            entry.configure(state="disabled")

    def select_file(self, key):
        path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if path:
            self.assign_file_to_key(key, path)

    def assign_file_to_key(self, key, path):
        """Asigna una ruta de archivo a una caja del formulario"""
        self.files_selected[key] = path
        entry = self.path_entries[key]
        entry.configure(state="normal")
        entry.delete(0, "end")
        entry.insert(0, os.path.basename(path))
        entry.configure(state="disabled")

    def on_file_drop(self, event):
        """Gestor de evento para cuando se sueltan uno o múltiples archivos en la aplicación"""
        files = self.tk.splitlist(event.data)
        if not files: return
        
        mode = self.mode_var.get()
        if mode == "Día Único":
            keys = ["Único"]
        elif mode == "Fin de Semana":
            keys = ["Viernes", "Sabado", "Domingo"]
        else:
            keys = ["Viernes", "Sabado", "Domingo", "Lunes"]
            
        pdf_files = [f for f in files if f.lower().endswith('.pdf')]
        
        # Asignamos en orden los archivos caídos a las muescas disponibles
        for i, f in enumerate(pdf_files):
            if i < len(keys):
                self.assign_file_to_key(keys[i], f)
        
        if len(pdf_files) > len(keys):
            self.log(f"⚠️ Soltaste {len(pdf_files)} PDFs pero se necesitan {len(keys)}. Los sobrantes se omitieron.")

    def log(self, msg):
        self.log_box.configure(state="normal")
        tag = "normal"
        if "❌" in msg or "Error" in msg or "Falló" in msg:
            tag = "error"
        elif "✅" in msg or "guardado" in msg or "generado" in msg:
            tag = "success"
        elif "⚠️" in msg:
            tag = "warning"
        elif "▶" in msg or "Iniciando" in msg or "---" in msg or "===" in msg:
            tag = "header"
            
        self.log_box.insert("end", msg + "\n", tag)
        self.log_box.see("end")
        self.log_box.configure(state="disabled")

    def poll_log_queue(self):
        try:
            while True:
                msg = self.log_queue.get_nowait()
                self.log(msg)
        except queue.Empty:
            pass
        self.after(100, self.poll_log_queue)

    def open_output_folder(self):
        if self.result_folder and os.path.exists(self.result_folder):
            open_folder(self.result_folder)

    def start_thread(self):
        if self.processing: return

        # Validación
        mode = self.mode_var.get()
        expected = 1 if mode == "Día Único" else (3 if mode == "Fin de Semana" else 4)
        if len(self.files_selected) < expected:
            messagebox.showwarning("Faltan Archivos", f"Por favor seleccione todos los archivos PDF requeridos para el modo '{mode}'.")
            return

        brand = self.brand_var.get()

        # Marcar que se limpiaran los shortcuts en la sesion a la primera iteracion
        clear_shortcuts = not self.folders_cleared_this_session
        self.folders_cleared_this_session = True

        # UI Lock
        self.processing = True
        self.btn_process.configure(state="disabled", text="Procesando...")
        self.btn_open_folder.configure(state="disabled")
        self.mode_selector.configure(state="disabled")
        self.brand_menu.configure(state="disabled")
        self.progress_bar.start()

        for btn in self.dynamic_fields_frame.winfo_children():
            if isinstance(btn, ctk.CTkButton):
                btn.configure(state="disabled")

        self.log_box.configure(state="normal")
        self.log_box.delete("1.0", "end")
        self.log_box.configure(state="disabled")

        # Setup Thread Logger
        logger = logging.getLogger("AppLogger")
        logger.setLevel(logging.INFO)
        if logger.handlers:
            logger.handlers = []
        queue_h = QueueHandler(self.log_queue)
        queue_h.setFormatter(logging.Formatter('%(message)s'))
        logger.addHandler(queue_h)

        files_locked = self.files_selected.copy()
        t = threading.Thread(target=self.run_process_wrapper, args=(brand, mode, files_locked, logger, clear_shortcuts))
        t.daemon = True
        t.start()

    def run_process_wrapper(self, brand, mode, files_dict, logger, clear_shortcuts):
        success, folder, report = run_processing_pipeline(brand, mode, files_dict, logger, clear_shortcuts)
        self.result_folder = os.path.join(os.path.expanduser("~/Desktop"), "guias-shein")
        self.after(0, lambda: self.finish_process(success, report))

    def finish_process(self, success, report):
        self.progress_bar.stop()
        self.processing = False
        self.btn_process.configure(state="normal", text="Procesar Guías")
        self.mode_selector.configure(state="normal")
        self.brand_menu.configure(state="normal")

        for btn in self.dynamic_fields_frame.winfo_children():
            if isinstance(btn, ctk.CTkButton):
                btn.configure(state="normal")

        accent = self.brand_colors.get(self.brand_var.get(), self.brand_colors["Marcas y Licencias"])["accent"]

        if success:
            self.btn_open_folder.configure(state="normal")
            
            # Construir Reporte
            report_str = "\n".join([f"• {dia}: {cnt} guías" for dia, cnt in report.items()])
            total_guias = sum(report.values())
            
            msg = f"Térmica copiadas a: /guias-shein/imprimir\nLáser copiada a: /guias-shein/correo\n\n📊 Total Guías: {total_guias}\n{report_str}"
            self.log(f"\n✅ CONTEO FINAL: {total_guias} Guías procesadas.")
            
            CustomAlert(self, "Proceso Terminado", msg, is_error=False, accent_color=accent)
        else:
            CustomAlert(self, "Error de Procesamiento", "Ocurrió un error. Revisa la consola de operaciones.", is_error=True, accent_color=accent)


if __name__ == "__main__":
    app = UnifiedSheinApp()
    app.mainloop()
